Now let me create the comprehensive backend with PDF parsing and LLM integration:
Action: file_editor create /app/backend/server.py --file-text "from fastapi import FastAPI, APIRouter, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional
import uuid
from datetime import datetime, timezone
import fitz  # PyMuPDF
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
from emergentintegrations.llm.chat import LlmChat, UserMessage


ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix=\"/api\")


# Define Models
class CoverLetterRequest(BaseModel):
    resume_text: str
    job_description: str

class CoverLetterResponse(BaseModel):
    cover_letter: str
    generated_at: str

class GeneratedCoverLetter(BaseModel):
    model_config = ConfigDict(extra=\"ignore\")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    resume_text: str
    job_description: str
    cover_letter: str
    generated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


# PDF Parsing Function
def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    \"\"\"Extract text from PDF using PyMuPDF\"\"\"
    try:
        doc = fitz.open(stream=pdf_bytes, filetype=\"pdf\")
        text = \"\"
        for page in doc:
            text += page.get_text()
        doc.close()
        return text.strip()
    except Exception as e:
        logging.error(f\"Error extracting PDF text: {str(e)}\")
        raise HTTPException(status_code=400, detail=f\"Failed to parse PDF: {str(e)}\")


# LLM Cover Letter Generation
async def generate_cover_letter_with_llm(resume_text: str, job_description: str) -> str:
    \"\"\"Generate a personalized cover letter using LLM\"\"\"
    
    # The \"Gold Standard\" Prompt - Acts as a Senior Recruiter
    system_message = \"\"\"You are a Senior Talent Acquisition Specialist and Professional Cover Letter Architect with 15+ years of experience at Fortune 500 companies and top-tier tech startups.

Your expertise:
- You've reviewed over 50,000 applications and know what makes candidates stand out
- You understand that recruiters spend 6-7 seconds on initial screening
- You spot \"fluff\" immediately and focus on quantifiable achievements
- You craft narratives that demonstrate cultural fit and genuine passion
- You write in a natural, engaging tone that feels human, not AI-generated

CRITICAL RULES:
1. NEVER start with generic phrases like \"I am writing to express my interest...\" or \"I am excited to apply...\"
2. Open with a HOOK - a compelling story, achievement, or insight related to the role
3. Use specific metrics and achievements from the resume
4. Show deep understanding of the company/role from the job description
5. Write conversationally but professionally - like a confident professional, not a desperate applicant
6. Keep it concise: 3-4 paragraphs, max 350 words
7. End with a memorable closing that reinforces value, not a generic \"thank you for your consideration\"

Structure:
- Opening Hook (2-3 sentences): Start with impact
- Body (2 paragraphs): Match achievements to job requirements with specifics
- Closing (2 sentences): Strong, confident finish

Remember: This letter should make the recruiter WANT to call this candidate immediately.\"\"\"

    try:
        # Initialize LLM Chat
        chat = LlmChat(
            api_key=os.environ.get('EMERGENT_LLM_KEY'),
            session_id=str(uuid.uuid4()),
            system_message=system_message
        ).with_model(\"openai\", \"gpt-4o\")
        
        # Create user prompt
        user_prompt = f\"\"\"Based on the following resume and job description, craft a world-class cover letter that will make this candidate stand out.

RESUME:
{resume_text}

JOB DESCRIPTION:
{job_description}

Generate a compelling, natural-sounding cover letter that:
1. Opens with an engaging hook (not generic opening)
2. Highlights 2-3 specific achievements that align with the job requirements
3. Demonstrates genuine understanding of the role and company
4. Uses a professional yet conversational tone
5. Ends with a strong, memorable closing
6. Stays under 350 words

Remember: Make it feel human, not AI-generated. Avoid clichés and generic phrases.\"\"\"

        user_message = UserMessage(text=user_prompt)
        
        # Get response from LLM
        response = await chat.send_message(user_message)
        
        return response.strip()
        
    except Exception as e:
        logging.error(f\"Error generating cover letter: {str(e)}\")
        raise HTTPException(status_code=500, detail=f\"Failed to generate cover letter: {str(e)}\")


# Generate DOCX file
def generate_docx(cover_letter_text: str) -> BytesIO:
    \"\"\"Generate a professional DOCX document\"\"\"
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add date
    date_para = doc.add_paragraph()
    date_para.add_run(datetime.now().strftime(\"%B %d, %Y\")).font.size = Pt(11)
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()  # Spacing
    
    # Add cover letter content
    paragraphs = cover_letter_text.split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            para = doc.add_paragraph(para_text.strip())
            para.style.font.name = 'Calibri'
            para.style.font.size = Pt(11)
    
    # Save to BytesIO
    docx_io = BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    return docx_io


# Generate PDF file
def generate_pdf(cover_letter_text: str) -> BytesIO:
    \"\"\"Generate a professional PDF document\"\"\"
    pdf_io = BytesIO()
    doc = SimpleDocTemplate(pdf_io, pagesize=letter,
                           topMargin=1*inch, bottomMargin=1*inch,
                           leftMargin=1*inch, rightMargin=1*inch)
    
    # Styles
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        spaceBefore=12,
        spaceAfter=12,
    )
    
    # Build content
    story = []
    
    # Add date
    date_text = datetime.now().strftime(\"%B %d, %Y\")
    story.append(Paragraph(date_text, styles['Normal']))
    story.append(Spacer(1, 0.3*inch))
    
    # Add paragraphs
    paragraphs = cover_letter_text.split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            story.append(Paragraph(para_text.strip(), body_style))
    
    doc.build(story)
    pdf_io.seek(0)
    return pdf_io


# API Routes
@api_router.get(\"/\")
async def root():
    return {\"message\": \"AI Cover Letter Architect API\", \"version\": \"1.0\"}


@api_router.post(\"/upload-resume\")
async def upload_resume(file: UploadFile = File(...)):
    \"\"\"Upload and parse PDF resume\"\"\"
    if not file.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail=\"Only PDF files are supported\")
    
    try:
        contents = await file.read()
        resume_text = extract_text_from_pdf(contents)
        
        if not resume_text:
            raise HTTPException(status_code=400, detail=\"Could not extract text from PDF\")
        
        return {
            \"success\": True,
            \"resume_text\": resume_text,
            \"filename\": file.filename
        }
    except Exception as e:
        logging.error(f\"Error uploading resume: {str(e)}\")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.post(\"/generate-cover-letter\", response_model=CoverLetterResponse)
async def generate_cover_letter(request: CoverLetterRequest):
    \"\"\"Generate a personalized cover letter using AI\"\"\"
    try:
        # Validate inputs
        if not request.resume_text or len(request.resume_text.strip()) < 50:
            raise HTTPException(status_code=400, detail=\"Resume text is too short or empty\")
        
        if not request.job_description or len(request.job_description.strip()) < 50:
            raise HTTPException(status_code=400, detail=\"Job description is too short or empty\")
        
        # Generate cover letter
        cover_letter = await generate_cover_letter_with_llm(
            request.resume_text,
            request.job_description
        )
        
        # Save to database
        cover_letter_doc = GeneratedCoverLetter(
            resume_text=request.resume_text[:1000],  # Store truncated version
            job_description=request.job_description[:1000],
            cover_letter=cover_letter
        )
        
        doc_dict = cover_letter_doc.model_dump()
        doc_dict['generated_at'] = doc_dict['generated_at'].isoformat()
        await db.cover_letters.insert_one(doc_dict)
        
        return CoverLetterResponse(
            cover_letter=cover_letter,
            generated_at=datetime.now(timezone.utc).isoformat()
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f\"Error generating cover letter: {str(e)}\")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.post(\"/download-docx\")
async def download_docx(request: dict):
    \"\"\"Download cover letter as DOCX\"\"\"
    try:
        cover_letter_text = request.get('cover_letter', '')
        if not cover_letter_text:
            raise HTTPException(status_code=400, detail=\"Cover letter text is required\")
        
        docx_file = generate_docx(cover_letter_text)
        
        return StreamingResponse(
            docx_file,
            media_type=\"application/vnd.openxmlformats-officedocument.wordprocessingml.document\",
            headers={\"Content-Disposition\": \"attachment; filename=cover_letter.docx\"}
        )
    except Exception as e:
        logging.error(f\"Error generating DOCX: {str(e)}\")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.post(\"/download-pdf\")
async def download_pdf(request: dict):
    \"\"\"Download cover letter as PDF\"\"\"
    try:
        cover_letter_text = request.get('cover_letter', '')
        if not cover_letter_text:
            raise HTTPException(status_code=400, detail=\"Cover letter text is required\")
        
        pdf_file = generate_pdf(cover_letter_text)
        
        return StreamingResponse(
            pdf_file,
            media_type=\"application/pdf\",
            headers={\"Content-Disposition\": \"attachment; filename=cover_letter.pdf\"}
        )
    except Exception as e:
        logging.error(f\"Error generating PDF: {str(e)}\")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.get(\"/history\")
async def get_history():
    \"\"\"Get generation history\"\"\"
    try:
        history = await db.cover_letters.find({}, {\"_id\": 0}).sort(\"generated_at\", -1).limit(10).to_list(10)
        
        for item in history:
            if isinstance(item.get('generated_at'), str):
                item['generated_at'] = datetime.fromisoformat(item['generated_at'])
        
        return {\"history\": history}
    except Exception as e:
        logging.error(f\"Error fetching history: {str(e)}\")
        raise HTTPException(status_code=500, detail=str(e))


# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=[\"*\"],
    allow_headers=[\"*\"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event(\"shutdown\")
async def shutdown_db_client():
    client.close()
"
Observation: Overwrite successful: /app/backend/server.py